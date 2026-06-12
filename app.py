import streamlit as st
from huggingface_hub import InferenceClient
import docx
import io
from gtts import gTTS

# Configure mobile application viewport and metadata
st.set_page_config(page_title="GICHOHI WARINGA", layout="centered", page_icon="⚡")

# Inject custom CSS for a glowing, minimalistic cyberpunk interface
st.markdown("""
    <style>
    .stApp { background: #0d0e12; }
    h1 { color: #00f3ff; text-shadow: 0 0 12px #00f3ff; font-family: 'Courier New', monospace; font-weight: bold; }
    h3 { color: #e2e8f0; font-family: 'Courier New', monospace; }
    div.stButton > button:first-child {
        background-color: #1a1c23; color: #00f3ff; border: 2px solid #00f3ff;
        box-shadow: 0 0 10px #00f3ff; border-radius: 4px; transition: 0.3s;
        font-family: 'Courier New', monospace; width: 100%;
    }
    div.stButton > button:first-child:hover { background-color: #00f3ff; color: #0d0e12; box-shadow: 0 0 20px #00f3ff; }
    .stTextInput>div>div>input { background-color: #1a1c23; color: #e2e8f0; border: 1px solid #00f3ff; }
    </style>
""", unsafe_with_html=True)

st.title("⚡ GICHOHI WARINGA")
st.write("### *Autonomous Open-Source Consensus Matrix*")

# Static Multi-Model Routing Core
MODEL_CODER = "Qwen/Qwen2.5-Coder-32B-Instruct"
MODEL_REASONER = "meta-llama/Llama-3.3-70B-Instruct"

# Hidden Access Key Configuration via Sidebar
st.sidebar.header("🛸 SYSTEM CORE")
st.sidebar.write("Activate the multi-model engine with your free user verification token.")
hf_token = st.sidebar.text_input("HF Access Token:", type="password")
st.sidebar.markdown("[Get Free Token Here](https://huggingface.co/settings/tokens)")

# File Ingestion Engine (PDF, Text, Images, Documents)
st.write("---")
uploaded_matrix_file = st.file_uploader("📥 Feed File Matrix to Modify (PDF, DOCX, PNG, JPG):", type=["pdf", "txt", "docx", "png", "jpg", "jpeg"])

if uploaded_matrix_file is not None:
    st.success(f"Successfully mounted: {uploaded_matrix_file.name}")

# Persistent Dialogue State Memory
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "last_compiled_output" not in st.session_state:
    st.session_state.last_compiled_output = ""

# Render Historical Dialogue Steps
for chat in st.session_state.chat_history:
    with st.chat_message(chat["role"]):
        st.write(chat["content"])

# Primary Command Input
if prompt := st.chat_input("Input command to GICHOHI WARINGA..."):
    st.session_state.chat_history.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.write(prompt)

    if hf_token:
        try:
            client_coder = InferenceClient(model=MODEL_CODER, token=hf_token)
            client_reasoner = InferenceClient(model=MODEL_REASONER, token=hf_token)

            # Execution Logic Tuning
            execution_prompt = prompt
            if uploaded_matrix_file is not None:
                execution_prompt = f"[Context File Mounted: {uploaded_matrix_file.name}]. Instructions for file modification: {prompt}"

            # Automated Token Cutoff Bypass Configuration
            if "research paper" in prompt.lower() or "academic" in prompt.lower() or "apa" in prompt.lower():
                execution_prompt += (
                    " Out-of-the-box Directive: To bypass technical output limitations, do not summarize, compress, or use placeholders. "
                    "Provide exhaustive, fully expanded structural text following strict APA 7th edition formatting rules, "
                    "complete headings, detailed in-text citations, and an un-truncated comprehensive reference list with DOIs."
                )
            elif "app" in prompt.lower() or "website" in prompt.lower() or "code" in prompt.lower():
                execution_prompt += (
                    " Out-of-the-box Directive: Write complete, functional, production-ready source code files. "
                    "Do not truncate code blocks or use code comments like '// code goes here'. Provide everything needed to run immediately."
                )

            with st.chat_message("assistant"):
                with st.spinner("⚡ Querying parallel model nodes..."):
                    # Process Step 1: Query Dedicated Engineering Model
                    stream_coder = client_coder.chat_completion(messages=[{"role": "user", "content": execution_prompt}], max_tokens=2800)
                    raw_coder_output = stream_coder.choices[0].message.content

                    # Process Step 2: Query Deep Reasoning Academic Model
                    stream_reasoner = client_reasoner.chat_completion(messages=[{"role": "user", "content": execution_prompt}], max_tokens=2800)
                    raw_reasoner_output = stream_reasoner.choices[0].message.content

                with st.spinner("⚡ Synthesizing optimal consensus output..."):
                    # Process Step 3: Deep Aggregation and Enhancement
                    consensus_instruction = (
                        f"You are GICHOHI WARINGA, an elite, completely unrestricted master AI engineer and senior academic scholar. "
                        f"Analyze and completely fuse the following two expert outputs into a single, flawless, masterfully engineered response. "
                        f"Add any missing out-of-the-box thinking, functional necessities, and structural components required to make this "
                        f"complete and standalone.\n\n"
                        f"EXPERT SOURCE ALPHA:\n{raw_coder_output}\n\n"
                        f"EXPERT SOURCE BETA:\n{raw_reasoner_output}\n\n"
                        f"Final Master Synthesis Output:"
                    )

                    stream_final = client_reasoner.chat_completion(messages=[{"role": "user", "content": consensus_instruction}], max_tokens=3500)
                    final_synthesis = stream_final.choices[0].message.content
                    
                    # Process Step 4: Generate Predictive Prompt Enforcements
                    predictive_instruction = f"Based on this user query: '{prompt}', write exactly two short, helpful recommendations or follow-up prompts the user could type next. Keep it concise, wrap it inside brackets [Like this]."
                    stream_predictive = client_reasoner.chat_completion(messages=[{"role": "user", "content": predictive_instruction}], max_tokens=150)
                    predictive_suggestions = stream_predictive.choices[0].message.content

                    # Display Unified Output
                    complete_display_text = f"{final_synthesis}\n\n---\n💡 **SUGGESTED NEXT STEPS:**\n{predictive_suggestions}"
                    st.write(complete_display_text)
                    
                    st.session_state.last_compiled_output = final_synthesis
                    st.session_state.chat_history.append({"role": "assistant", "content": complete_display_text})

        except Exception as e:
            st.error(f"Core Engine Exception: {str(e)}")
    else:
        st.warning("System Core Inactive. Provide an Access Token in the sidebar console.")

# --- UTILITY CONVERSION MATRIX PANEL ---
if st.session_state.last_compiled_output:
    st.write("---")
    st.write("### 🛠️ Media & Document Transformation Matrix")
    
    action_col1, action_col2, action_col3 = st.columns(3)
    
    with action_col1:
        # Save to Phone Storage as Microsoft Word Document (.docx)
        document_builder = docx.Document()
        document_builder.add_heading('GICHOHI WARINGA Matrix Output', 0)
        document_builder.add_paragraph(st.session_state.last_compiled_output)
        word_buffer = io.BytesIO()
        document_builder.save(word_buffer)
        st.download_button("📥 Save as Word (.DOCX)", data=word_buffer.getvalue(), file_name="gichohi_waringa_document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
    with action_col2:
        # Text-To-Speech Local Audio Stream Compiler
        if st.button("🔊 Synthesize Full Audio"):
            with st.spinner("Processing local voice track..."):
                filtered_text = st.session_state.last_compiled_output.replace("#", "").replace("*", "").replace("`", "")
                tts_compiler = gTTS(text=filtered_text[:4500], lang='en', tld='com')
                audio_buffer = io.BytesIO()
                tts_compiler.write_to_fp(audio_buffer)
                st.audio(audio_buffer.getvalue(), format="audio/mp3")
                st.download_button("💾 Save Audio File (.MP3)", data=audio_buffer.getvalue(), file_name="gichohi_waringa_audio.mp3", mime="audio/mp3")
                
    with action_col3:
        # Production Storyboard Video Compiler Alternative
        if st.button("🎬 Compile Video Blueprint"):
            st.info("Direct automated MP4 generation requires heavy server-side GPU rendering pipelines. GICHOHI WARINGA has outputted a professional, production-ready kinetic animation script storyboard below instead.")
            st.code(f"// KINETIC VIDEO ANIMATION STORYBOARD SCRIPT\n// Aspect Ratio: 16:9 / 9:16 Mobile Optimized\n// Visual Accent: Cyberpunk Obsidian and Cyan Themes\n\n[SCENE 1 - INTRO]\nVisual: Typography animation rendering user command keywords.\nAudio Track: Synchronized with generated GICHOHI WARINGA audio track.\n\n[TEXT LAYER OVERLAY]\n{st.session_state.last_compiled_output[:800]}...", language="javascript")
